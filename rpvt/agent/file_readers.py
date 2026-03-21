"""File readers for different formats.

Extracts text content from various file types so the agent
can ingest them into KV memory.
"""

import csv
import io
import json
from pathlib import Path


def read_file(path):
    """Read a file and return its text content.

    Supports: .txt, .md, .csv, .json, .jsonld,
              .pdf, .xlsx, .docx, .py, .html
    Returns (text, error) tuple.
    """
    path = Path(path)
    suffix = path.suffix.lower()

    try:
        if suffix in (".txt", ".md", ".py", ".html", ".htm", ".xml", ".log"):
            return path.read_text(encoding="utf-8", errors="replace"), None

        elif suffix == ".csv":
            return _read_csv(path), None

        elif suffix in (".json", ".jsonld"):
            return _read_json(path), None

        elif suffix == ".pdf":
            return _read_pdf(path), None

        elif suffix == ".xlsx":
            return _read_xlsx(path), None

        elif suffix == ".docx":
            return _read_docx(path), None

        elif suffix == ".zip":
            return _read_zip(path), None

        elif suffix == ".pdb":
            return path.read_text(encoding="utf-8", errors="replace"), None

        else:
            # Try as text
            try:
                return path.read_text(encoding="utf-8", errors="replace"), None
            except Exception:
                return None, f"Unsupported file type: {suffix}"

    except Exception as e:
        return None, f"Error reading {path.name}: {e}"


def _read_csv(path):
    text = path.read_text(encoding="utf-8", errors="replace")
    reader = csv.reader(io.StringIO(text))
    rows = list(reader)
    if not rows:
        return ""
    # Format as readable table
    lines = []
    for i, row in enumerate(rows):
        lines.append(" | ".join(str(c) for c in row))
    return "\n".join(lines)


def _read_json(path):
    with open(path, "r", encoding="utf-8") as f:
        data = json.load(f)
    return json.dumps(data, indent=2, ensure_ascii=False)


def _read_pdf(path):
    import fitz  # pymupdf
    doc = fitz.open(str(path))
    pages = []
    for page in doc:
        pages.append(page.get_text())
    doc.close()
    return "\n\n".join(pages)


def _read_xlsx(path):
    import openpyxl
    wb = openpyxl.load_workbook(str(path), data_only=True)
    sheets = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        lines = [f"Sheet: {sheet_name}"]
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) if c is not None else "" for c in row]
            lines.append(" | ".join(cells))
        sheets.append("\n".join(lines))
    wb.close()
    return "\n\n".join(sheets)


def _read_docx(path):
    import docx
    doc = docx.Document(str(path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # Also read tables
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            paragraphs.append(" | ".join(cells))
    return "\n".join(paragraphs)


def _read_zip(path):
    import zipfile
    import tempfile
    import os

    texts = []
    with zipfile.ZipFile(str(path), "r") as zf:
        with tempfile.TemporaryDirectory() as tmpdir:
            zf.extractall(tmpdir)
            for root, dirs, files in os.walk(tmpdir):
                for fname in sorted(files):
                    fpath = os.path.join(root, fname)
                    rel = os.path.relpath(fpath, tmpdir)
                    content, err = read_file(fpath)
                    if content:
                        texts.append(f"=== File: {rel} ===\n{content}")
                    elif err:
                        texts.append(f"=== File: {rel} === (error: {err})")
    return "\n\n".join(texts)
